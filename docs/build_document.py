from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
BUILD = DOCS / "build_assets"
OUT = DOCS / "Growise_SmartReco_Project_Documentation.docx"
BUILD.mkdir(exist_ok=True)

INK = "#18202B"
MUTED = "#5D6675"
PURPLE = "#5A47DC"
PURPLE_SOFT = "#EEEAFE"
PURPLE_DARK = "#332686"
MINT = "#E4F7EF"
GREEN = "#167A52"
PAPER = "#FAFBFD"
LINE = "#DDE2EA"
SLATE = "#EFF2F7"
WHITE = "#FFFFFF"


def font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
        "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size, index=0)
    return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline=None, radius=22, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw, box, text, fnt, fill, max_chars=30, line_gap=8):
    left, top, right, bottom = box
    lines = []
    for paragraph in text.split("\n"):
        lines.extend(wrap(paragraph, width=max_chars) or [""])
    line_heights = [draw.textbbox((0, 0), line, font=fnt)[3] for line in lines]
    total = sum(line_heights) + line_gap * (len(lines) - 1)
    y = top + ((bottom - top) - total) / 2
    for line, h in zip(lines, line_heights):
        bbox = draw.textbbox((0, 0), line, font=fnt)
        x = left + ((right - left) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=fnt, fill=fill)
        y += h + line_gap


def arrow(draw, start, end, fill=PURPLE, width=7):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - 22 * sign, y2 - 13), (x2 - 22 * sign, y2 + 13)]
    else:
        sign = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 13, y2 - 22 * sign), (x2 + 13, y2 - 22 * sign)]
    draw.polygon(points, fill=fill)


def diagram_architecture(path):
    img = Image.new("RGB", (1800, 980), PAPER)
    draw = ImageDraw.Draw(img)
    h1, h2, body = font(40, True), font(27, True), font(22)
    draw.text((75, 42), "Growise system architecture", font=h1, fill=INK)
    draw.text((75, 98), "A split web experience with a backend decision layer and durable operational evidence.", font=body, fill=MUTED)
    boxes = [
        ((80, 250, 405, 465), "Learner & Admin UI\nNext.js App Router", PURPLE_SOFT, PURPLE_DARK),
        ((520, 250, 845, 465), "FastAPI JSON API\nAuth · products · events", "#E8F2FF", "#155E9C"),
        ((960, 165, 1300, 380), "SQLite / Postgres\nUsers · catalog · events\nrecommendations", "#F7EFE2", "#805A13"),
        ((960, 515, 1300, 730), "Chroma vector store\nEmbeddings + metadata\nsemantic retrieval", MINT, GREEN),
        ((1415, 335, 1735, 570), "LangGraph agent\nSignal → retrieve →\nrerank → deliver", "#F3EEFF", PURPLE_DARK),
    ]
    for box, label, fill, stroke in boxes:
        rounded_box(draw, box, fill, stroke)
        center_text(draw, box, label, h2, INK, max_chars=21)
    arrow(draw, (405, 357), (520, 357))
    arrow(draw, (845, 310), (960, 272), fill="#155E9C")
    arrow(draw, (845, 400), (960, 622), fill=GREEN)
    arrow(draw, (1300, 272), (1415, 395), fill=PURPLE)
    arrow(draw, (1300, 622), (1415, 510), fill=PURPLE)
    arrow(draw, (1415, 440), (845, 440), fill=PURPLE_DARK)
    draw.text((620, 490), "personalised recommendation + audit trace", font=body, fill=PURPLE_DARK)
    img.save(path)


def diagram_agent(path):
    img = Image.new("RGB", (1800, 840), PAPER)
    draw = ImageDraw.Draw(img)
    h1, node, body = font(40, True), font(25, True), font(20)
    draw.text((75, 45), "Recommendation decision loop", font=h1, fill=INK)
    draw.text((75, 102), "The agent reasons from consented behavioural evidence and is grounded in the live catalog.", font=body, fill=MUTED)
    steps = [
        (70, "1. Observe", "Views, searches, clicks, qualified dwell"),
        (405, "2. Profile", "Interest points, recency, exclusions"),
        (740, "3. Retrieve", "Top catalog candidates from Chroma"),
        (1075, "4. Rerank", "Mesh selects fit + path variety"),
        (1410, "5. Deliver", "Narrative + ranked courses, persisted"),
    ]
    for i, (x, title, detail) in enumerate(steps):
        box = (x, 305, x + 280, 555)
        fill = PURPLE_SOFT if i in (1, 3) else WHITE
        rounded_box(draw, box, fill, PURPLE if i in (1, 3) else LINE)
        center_text(draw, (x + 20, 335, x + 260, 435), title, node, INK, max_chars=20)
        center_text(draw, (x + 25, 435, x + 255, 520), detail, body, MUTED, max_chars=22)
        if i < len(steps) - 1:
            arrow(draw, (x + 280, 430), (x + 335, 430), fill=PURPLE)
    draw.rounded_rectangle((465, 655, 1335, 745), radius=18, fill=MINT, outline="#B9E8D4", width=2)
    center_text(draw, (490, 670, 1310, 730), "Efficiency gate: serve the cached recommendation until there is meaningful new signal and the cooldown has elapsed; a deliberate refresh may bypass the gate but still requires activity.", body, GREEN, max_chars=96)
    img.save(path)


def diagram_data(path):
    img = Image.new("RGB", (1800, 900), PAPER)
    draw = ImageDraw.Draw(img)
    h1, node, body = font(40, True), font(25, True), font(20)
    draw.text((75, 45), "Catalog dual-write and recommendation persistence", font=h1, fill=INK)
    draw.text((75, 102), "The catalog remains visible and auditable even if semantic sync needs a retry.", font=body, fill=MUTED)
    left = (110, 270, 530, 570)
    mid_top = (700, 185, 1110, 410)
    mid_bottom = (700, 505, 1110, 730)
    right = (1280, 270, 1690, 570)
    for box, label, fill, stroke in [
        (left, "Admin creates / edits course\nTitle, description, category, price, tags", PURPLE_SOFT, PURPLE_DARK),
        (mid_top, "Primary database\nProduct record committed first\nvector_synced status retained", "#E8F2FF", "#155E9C"),
        (mid_bottom, "Chroma collection\nEmbedding + searchable metadata\nupsert / delete", MINT, GREEN),
        (right, "Recommendation records\nActive item set + reasons\nrun + Mesh telemetry", "#F7EFE2", "#805A13"),
    ]:
        rounded_box(draw, box, fill, stroke)
        center_text(draw, box, label, node, INK, max_chars=26)
    arrow(draw, (530, 350), (700, 300), fill="#155E9C")
    arrow(draw, (530, 495), (700, 620), fill=GREEN)
    arrow(draw, (1110, 300), (1280, 365), fill=PURPLE)
    arrow(draw, (1110, 620), (1280, 475), fill=PURPLE)
    draw.rounded_rectangle((640, 785, 1160, 850), radius=14, fill=WHITE, outline=LINE, width=2)
    center_text(draw, (660, 795, 1140, 840), "If vector sync fails: SQL record remains, state becomes unsynced, and Catalog health exposes a safe retry.", body, MUTED, max_chars=72)
    img.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.lstrip("#"))
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcw = cell._tc.tcPr.first_child_found_in("w:tcW")
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                cell._tc.tcPr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_run_font(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    run.bold = bold
    run.italic = italic


def add_para(doc, text="", size=10.5, color=INK, bold=False, italic=False, align=None, before=0, after=7, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.25
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def add_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    set_run_font(r, size=8.5, color=PURPLE, bold=True)
    r.font.all_caps = True
    return p


def add_caption(doc, label, action):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(label + " ")
    set_run_font(r, size=8.5, color=MUTED, italic=True)
    r2 = p.add_run("Reviewer action: ")
    set_run_font(r2, size=8.5, color=PURPLE_DARK, bold=True)
    r3 = p.add_run(action)
    set_run_font(r3, size=8.5, color=MUTED)
    return p


def add_figure(doc, image_path, label, action, width=6.05):
    doc.add_picture(str(image_path), width=Inches(width))
    add_caption(doc, label, action)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, SLATE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9, color=PURPLE_DARK, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(value)
            set_run_font(r, size=8.9, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=8.5, color=MUTED)


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK.lstrip("#"))
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in [(1, 18, PURPLE_DARK, 16, 7), (2, 13, PURPLE, 12, 5), (3, 11, PURPLE_DARK, 8, 4)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("GROWISE  |  SMARTRECO PROJECT DOCUMENTATION"), size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    set_run_font(footer.add_run("Growise · SmartReco Hackathon · Page "), size=8.5, color=MUTED)
    add_page_field(footer)


def preview(src_name, out_name, max_height):
    with Image.open(ASSETS / src_name) as image:
        image = image.convert("RGB")
        crop = image.crop((0, 0, image.width, min(image.height, max_height)))
        crop.save(BUILD / out_name, quality=92)
    return BUILD / out_name


def main():
    architecture = BUILD / "architecture.png"
    agent_flow = BUILD / "agent_flow.png"
    data_flow = BUILD / "data_flow.png"
    diagram_architecture(architecture)
    diagram_agent(agent_flow)
    diagram_data(data_flow)

    screenshots = {
        "login": preview("01-user-login.png", "login.jpg", 903),
        "signup": preview("02-signup.png", "signup.jpg", 764),
        "landing": preview("03-learner-landing.png", "landing.jpg", 1220),
        "catalog": preview("04-course-catalog.png", "catalog.jpg", 720),
        "detail": preview("05-course-detail.png", "detail.jpg", 1050),
        "admin_login": preview("06-admin-login.png", "admin_login.jpg", 903),
        "ops": preview("07-admin-agent-ops.png", "ops.jpg", 960),
        "events": preview("08-admin-events.png", "events.jpg", 950),
        "traces": preview("09-admin-trace-explorer.png", "traces.jpg", 1100),
        "health": preview("10-admin-catalog-health.png", "health.jpg", 720),
        "manage": preview("11-admin-course-management.png", "manage.jpg", 869),
        "for_you": preview("12-for-you.png", "for_you.jpg", 1050),
        "learning": preview("13-my-learning.png", "learning.jpg", 1050),
    }

    doc = Document()
    configure_doc(doc)

    # Cover: editorial_cover variant, resolved to Growise brand tokens.
    for _ in range(5):
        add_para(doc, "", after=0)
    add_para(doc, "SMARTRECO HACKATHON · PRODUCT & ARCHITECTURE DOCUMENTATION", size=9, color=PURPLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_para(doc, "Growise", size=34, color=PURPLE_DARK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Behavioural AI recommendations for an online learning marketplace", size=17, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_para(doc, "A product walkthrough, implementation architecture, and reviewer guide", size=11.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=45)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    for cell, label, value in zip(table.rows[0].cells, ["EXPERIENCE", "RECOMMENDATIONS", "OPERATIONS"], ["Learner journeys", "Grounded, personal", "Auditable admin console"]):
        set_cell_shading(cell, PURPLE_SOFT)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        set_run_font(p.add_run(label + "\n"), size=8.5, color=PURPLE, bold=True)
        set_run_font(p.add_run(value), size=10.5, color=INK, bold=True)
    add_para(doc, "", after=44)
    add_para(doc, "Prepared from the live local application experience and the current source implementation.", size=9.5, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()

    add_label(doc, "01 · Introduction")
    add_heading(doc, "What Growise is", 1)
    add_para(doc, "Growise is a course marketplace that replaces generic “related courses” with a behavioural AI recommendation loop. It observes the meaningful choices a learner makes—searches, course views, card clicks, qualified dwell time and enrollment intent—and converts those signals into a concise, catalog-grounded next-step recommendation.")
    add_para(doc, "The product is designed as two complementary experiences. Learners receive a focused route through a large course catalog, while administrators can manage catalog quality and inspect the evidence, agent steps, model usage and outcomes behind every recommendation.")
    add_heading(doc, "The problem", 1)
    add_para(doc, "Traditional course discovery expects learners to articulate their goal, filter a large catalog and repeatedly evaluate alternatives. That creates three recurring problems: choice overload, recommendations that are not connected to current intent, and opaque systems that give operators no credible way to explain why a course was selected.")
    add_heading(doc, "What we are solving", 1)
    add_para(doc, "Growise turns passive browsing into a lightweight learning signal. It makes a recommendation only when evidence exists, retrieves candidate courses from the real catalog, applies relevance and diversity checks, generates a tailored explanation through Mesh, and stores the result with an inspectable trace. The result is a more useful learner experience without treating every click as a reason to call an LLM.")
    add_table(doc, ["Design goal", "How Growise addresses it"], [
        ("Personal relevance", "Weighted interest signals, recency and exclusions shape a learner profile before retrieval."),
        ("Catalog grounding", "Chroma retrieves real course candidates; generated product IDs are validated against the candidate set."),
        ("Cost discipline", "A trigger gate serves the cached recommendation until there is enough new meaningful activity and the cooldown is complete."),
        ("Operational trust", "The admin console exposes events, recommendation traces, Mesh calls and vector-sync health."),
    ], [2400, 6960])

    add_label(doc, "02 · Architecture")
    add_heading(doc, "Complete architecture", 1)
    add_para(doc, "The frontend and backend are intentionally separated. Next.js provides server-rendered pages and interactive client components. FastAPI owns authentication, catalog mutations, behavioural event intake, recommendation delivery, operations data and the agent workflow. JWT bearer tokens keep cross-origin development simple; FastAPI enforces user and administrator boundaries on protected endpoints.")
    add_figure(doc, architecture, "Figure 1. End-to-end Growise architecture.", "Read left to right: the web UI records and requests; FastAPI coordinates durable data and the agent; the vector store supplies grounded candidates.")
    add_heading(doc, "Component responsibilities", 2)
    add_table(doc, ["Layer", "Responsibility", "Key implementation choices"], [
        ("Next.js 16 frontend", "Learner and admin journeys, course discovery, auth UI and event capture.", "App Router, Tailwind UI, non-blocking tracker, API client."),
        ("FastAPI backend", "JSON API, JWT validation, administrative guards and orchestration.", "Routers for auth, products, events, enrollments, recommendations and agent ops."),
        ("SQL store", "System of record for users, courses, events, enrollments, recommendations and run telemetry.", "SQLAlchemy entities and durable recommendation history."),
        ("Chroma", "Semantic catalog retrieval.", "Course metadata + local all-MiniLM-L6-v2 embeddings; explicit sync status."),
        ("LangGraph + Mesh", "Structured recommendation reasoning and persuasive narrative generation.", "Bounded workflow, relevance checks, candidate validation and measured usage."),
    ], [1900, 3500, 3960])

    add_heading(doc, "Behavioural intelligence and agent workflow", 1)
    add_para(doc, "The tracker queues events in memory and flushes them every five seconds or at twenty events. It uses keepalive on page hide/unload so a final batch can complete without blocking navigation. Short visits remain available to operations, but only dwell at or above the qualifying threshold affects recommendation reasoning.")
    add_figure(doc, agent_flow, "Figure 2. Behavioural signal to delivered recommendation.", "Use this flow to verify the important boundary: an LLM never invents from an unconstrained catalog; it ranks and explains retrieved candidates.")
    add_para(doc, "In the profile stage, adjacent card-click and product-view events are collapsed into a single course-opening session. Interest is weighted by event type, dwell, enrollment intent and recency. Enrolled courses and recently dismissed recommendations become exclusions. Retrieval obtains a small candidate set from Chroma, relevance can be checked and refined once, and Mesh then reranks for fit and learning-path variety before the final narrative and ranked items are persisted.")

    add_heading(doc, "Catalog integrity and persistence", 1)
    add_para(doc, "Course administration uses a deliberate dual-write pattern. A product is committed to the primary database before Chroma is updated. If semantic upsert fails, the course remains visible with vector_synced=false and the failure is retained for the catalog-health view. This avoids losing a valid catalog edit while making retrieval readiness explicit and safely retryable.")
    add_figure(doc, data_flow, "Figure 3. Dual-write catalog flow and recommendation evidence.", "Verify that the primary record is committed first, then inspect Catalog health to see whether the semantic copy is ready or needs a retry.")

    add_label(doc, "03 · Data, security and efficiency")
    add_heading(doc, "Data model and lifecycle", 1)
    add_table(doc, ["Record", "Purpose", "Lifecycle / guardrail"], [
        ("users", "Email/password credential, role and tracking preference.", "Password hash; JWT bearer token; user or admin role."),
        ("products", "Course catalogue and vector-sync state.", "Written first to SQL; semantic copy upserted; failed sync is recoverable."),
        ("events", "Meaningful learner actions and metadata.", "Batched async intake; short dwell retained but not always agent-eligible."),
        ("recommendations + items", "Narrative and ranked, real-course selections.", "New result becomes active; previous result is retained as history."),
        ("agent_runs + steps + Mesh logs", "Traceable decision and model-use evidence.", "Admin-only review surface; measured latency and token metadata."),
    ], [2100, 3500, 3760])
    add_heading(doc, "Privacy and production guardrails", 2)
    add_para(doc, "Tracking consent is explicit at sign-up. A learner may opt out, which stops client-side event queuing and prevents recommendation triggering without preventing ordinary use of the marketplace. The collector avoids raw scroll and mousemove streams; it records meaningful actions and one computed time-on-page event instead. Admin views are role-gated, and learner emails shown in operations are confined to that admin-only interface.")
    add_heading(doc, "LLM-call efficiency", 2)
    add_para(doc, "The recommendation endpoint first asks whether regeneration is warranted. No signal means no agent run. If an active recommendation exists, the backend requires enough new meaningful events since that recommendation and an elapsed cooldown; otherwise it returns the cached result. The explicit Refresh action is a deliberate user request, but it still requires at least some activity. This keeps the system responsive, reduces duplicate calls and preserves a stable recommendation until the learner’s behaviour materially changes.")

    doc.add_page_break()
    add_label(doc, "04 · Learner experience")
    add_heading(doc, "Learner journey and annotated screenshots", 1)
    add_para(doc, "The learner experience is a continuous discovery loop: understand the value, create an account with consent, explore the catalog, inspect a course and return to a tailored next step. Each capture below includes a concrete reviewer action.")
    add_figure(doc, screenshots["landing"], "Figure 4. Learner landing page.", "Select “Explore courses” to enter discovery, or use “See how it works” to connect the three-step story with the product behaviour.")
    doc.add_page_break()
    add_figure(doc, screenshots["signup"], "Figure 5. Sign-up page with tracking consent.", "Create an account with the consent switch enabled to allow the recommendation agent to learn from meaningful browsing; switch it off to use Growise without recommendations.")
    doc.add_page_break()
    add_figure(doc, screenshots["login"], "Figure 6. Learner login.", "Use “Continue as guest” for the seeded learner walkthrough, or sign in with an existing email and password to restore courses and recommendations.")
    doc.add_page_break()
    add_figure(doc, screenshots["catalog"], "Figure 7. Filtered course catalogue.", "Choose a category, level or price filter and open a course card; the resulting views and clicks become lightweight recommendation evidence after sign-in.")
    doc.add_page_break()
    add_figure(doc, screenshots["detail"], "Figure 8. Course detail page.", "Review the outcomes and expandable curriculum, then use “Enroll now” when ready; this is also the place where a signed-in learner sees the current learning signal.")
    doc.add_page_break()
    add_figure(doc, screenshots["for_you"], "Figure 9. Personalised “For you” page.", "Review the narrative and “Why this fits” explanations, then choose a recommended course or use “Not for me” to exclude a poor fit from future recommendations.")
    doc.add_page_break()
    add_figure(doc, screenshots["learning"], "Figure 10. My learning library.", "Open an enrolled course to continue learning; courses already owned are excluded from future recommendation candidates.")

    doc.add_page_break()
    add_label(doc, "05 · Administrator experience")
    add_heading(doc, "Admin journey and annotated screenshots", 1)
    add_para(doc, "The admin console makes the recommendation system inspectable rather than magical. It is where a reviewer can confirm the operating health of the catalog, inspect consented signals and audit the exact decision path behind a learner-facing recommendation.")
    add_figure(doc, screenshots["admin_login"], "Figure 11. Administrator login path.", "Choose “Continue as admin” on the sign-in screen to enter the seeded operational workspace and land on the recommendation dashboard.")
    doc.add_page_break()
    add_figure(doc, screenshots["ops"], "Figure 12. Recommendation operations dashboard.", "Start here to review daily run volume, Mesh tokens, latency and recommendation engagement; open a recent run for its underlying trace.")
    doc.add_page_break()
    add_figure(doc, screenshots["events"], "Figure 13. Live consented event stream.", "Filter by event type, learner, course or search query to check the behavioural evidence available to the system.")
    doc.add_page_break()
    add_figure(doc, screenshots["traces"], "Figure 14. Agent trace explorer.", "Select a recommendation run to verify the learner signal, retrieved candidates, exclusions, reranking, delivered courses and measured Mesh calls.")
    doc.add_page_break()
    add_figure(doc, screenshots["health"], "Figure 15. Catalog health.", "Confirm retrieval coverage is 100%; if a course is unsynced, use “Retry unsynced courses” to repair only the affected vector records.")
    doc.add_page_break()
    add_figure(doc, screenshots["manage"], "Figure 16. Course management.", "Use “Add course” or “Edit” to maintain the catalog; check the Synced column after each change to confirm semantic retrieval readiness.")

    doc.add_page_break()
    add_label(doc, "06 · Reviewer quick path")
    add_heading(doc, "Suggested end-to-end review", 1)
    add_para(doc, "A concise demo sequence proves both the learner value and the operational controls. First, use the landing page to explain the promise. Next, sign in as the seeded learner, browse courses in one theme, open a few details and view the “For you” page. Finally, sign in as the administrator and use Agent Ops, Events, Trace Explorer and Catalog health to show that the recommendation is grounded, observable and operationally safe.")
    add_table(doc, ["Step", "Reviewer action", "Expected evidence"], [
        ("1", "Open the landing page and catalogue.", "Clear value proposition, filters and real catalog."),
        ("2", "Sign in as guest and explore several courses.", "Consented views, clicks and dwell data appear in the live event stream."),
        ("3", "Open “For you” or choose Update path.", "A short personalised narrative and real course recommendations."),
        ("4", "Sign in as admin and open the latest trace.", "Signal, retrieval candidates, exclusions, ranking, recommendation and Mesh metrics."),
        ("5", "Open Catalog health and Course management.", "Dual-write health, safe retry control and course-level sync status."),
    ], [700, 4200, 4460])
    add_heading(doc, "What makes this submission complete", 2)
    add_para(doc, "Growise delivers the SmartReco brief as a working platform rather than a static recommendation widget: authentication and role separation, course CRUD with durable vector sync state, efficient behavioural tracking, a structured LangGraph workflow through Mesh, catalog-grounded personalised copy, stored recommendation history and a full administrator audit surface. The screenshots in this document are captured from the live local application and map directly to reviewer actions.")

    doc.core_properties.title = "Growise | SmartReco Project Documentation"
    doc.core_properties.subject = "Product walkthrough, architecture, and reviewer guide"
    doc.core_properties.author = "Growise"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
